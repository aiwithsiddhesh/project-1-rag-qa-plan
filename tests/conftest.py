import os
import sys
from pathlib import Path
from unittest.mock import MagicMock

import numpy as np
import pytest
from langchain_core.documents import Document
from langchain_core.messages import AIMessage

# Prevent the src.config module-level Settings() singleton from crashing on import
# when OPENAI_API_KEY is not present in the local environment.
os.environ.setdefault("OPENAI_API_KEY", "test-api-key")


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


@pytest.fixture
def sample_txt_path(tmp_path: Path) -> Path:
    path = tmp_path / "sample.txt"
    path.write_text(
        "First paragraph with enough text to form a meaningful unit of content.\n\n"
        "Second paragraph with additional content for testing purposes here.\n\n"
        "Third paragraph providing yet more content for test completeness overall."
    )
    return path


@pytest.fixture
def sample_pdf_path(tmp_path: Path) -> Path:
    fitz = pytest.importorskip(
        "fitz", reason="PyMuPDF not available on this system", exc_type=ImportError
    )
    path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (50, 50),
        "PDF first paragraph with enough content for loading tests.\n\n"
        "PDF second paragraph with additional content for verification.\n\n"
        "PDF third paragraph rounding out the test document content.",
    )
    doc.save(str(path))
    doc.close()
    return path


@pytest.fixture
def sample_docx_path(tmp_path: Path) -> Path:
    from docx import Document as DocxDocument

    path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("DOCX first paragraph with sufficient content for loading tests.")
    doc.add_paragraph("DOCX second paragraph providing additional content for tests.")
    doc.save(str(path))
    return path


@pytest.fixture
def sample_docs_dir(
    tmp_path: Path,
    sample_txt_path: Path,  # noqa: ARG001
    sample_pdf_path: Path,  # noqa: ARG001
) -> Path:
    return tmp_path


@pytest.fixture
def empty_docs_dir(tmp_path: Path) -> Path:
    docs_dir = tmp_path / "empty"
    docs_dir.mkdir()
    return docs_dir


@pytest.fixture
def sample_chunks() -> list[Document]:
    return [
        Document(
            page_content=f"Sample chunk {i} with content for testing retrieval pipelines.",
            metadata={
                "source_file": f"doc_{i % 2}.txt",
                "file_type": "txt",
                "chunk_index": i,
                "total_chunks": 5,
            },
        )
        for i in range(5)
    ]


@pytest.fixture
def mock_embedding_model() -> MagicMock:
    model = MagicMock()
    vector = np.ones(384).tolist()
    model.embed_documents.side_effect = lambda texts: [vector for _ in texts]
    model.embed_query.return_value = vector
    return model


@pytest.fixture
def mock_llm() -> MagicMock:
    llm = MagicMock()
    llm.invoke.return_value = AIMessage(content="This is a mock hypothetical answer.")
    return llm


@pytest.fixture
def mock_vectorstore() -> MagicMock:
    store = MagicMock()
    fake_docs = [
        Document(
            page_content=f"Result chunk {i}",
            metadata={
                "source_file": f"result_doc_{i}.txt",
                "chunk_index": i,
                "total_chunks": 3,
            },
        )
        for i in range(3)
    ]
    store.similarity_search_with_score.return_value = [
        (doc, 0.9 - i * 0.1) for i, doc in enumerate(fake_docs)
    ]
    store.max_marginal_relevance_search_with_score_by_vector.return_value = [
        (doc, 0.9 - i * 0.1) for i, doc in enumerate(fake_docs)
    ]
    return store
